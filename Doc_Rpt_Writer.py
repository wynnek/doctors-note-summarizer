"""
This Python script uses the Faker, pandas, docx, nltk, and streamlit libraries to generate, process, and
summarize fake doctor's reports.

Here's a summary of what the script does:

1. **Imports necessary libraries**: The script begins by importing the necessary Python libraries. Faker 
is used to generate fake data, pandas is used for data manipulation, docx is used to read and write Word 
documents, nltk is used for natural language processing, and streamlit is used to create a web app.

2. **Initializes Faker**: Faker is a Python package that generates fake data. In this script, it's used to 
generate fake doctor's reports.

3. **Generates a fake doctor's report and saves it to a text file, Word file, and Excel file**: The script 
uses Faker to generate a fake doctor's report, which it then writes to a text file, Word document, and Excel file.

4. **Downloads necessary NLTK data**: The script uses the nltk library to download the 'punkt' and 'stopwords' 
datasets, which are used for tokenizing text and filtering out common words, respectively.

5. **Loads and reads files**: The script includes functions to load and read text files, Word documents, and 
Excel files. These functions are used to read the doctor's reports that were previously saved.

6. **Summarizes reports**: The script includes a function to summarize a report. This function tokenizes the 
text, filters out common words, and identifies the most important sentences to include in the summary.

7. **Writes summaries to files**: The script includes functions to write the summaries to text files, Word 
documents, and Excel files.

8. **Creates a Streamlit app**: The script uses Streamlit to create a web app that allows users to upload 
a doctor's report in .txt, .docx, or .xlsx format. The app then reads the report, generates a summary, and 
writes the summary to a .txt, .docx, or .xlsx file.
"""

import streamlit as st
from faker import Faker
import pandas as pd
import tempfile
from docx import Document
from openpyxl import load_workbook
import os
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize

# Initialize Faker
fake = Faker()

# Generate a fake doctor's report and save it to a text file
with open('doctors_report.txt', 'w') as f:
    f.write(fake.text())

# Generate a fake doctor's report and save it to a Word file
doc = Document()
doc.add_paragraph(fake.text())
doc.save('doctors_report.docx')

# Generate a fake doctor's report and save it to an Excel file
df = pd.DataFrame({'Report': [fake.text()]})
df.to_excel('doctors_report.xlsx', index=False)

# Check if the necessary NLTK packages are downloaded
def check_nltk_packages():
    nltk_packages = ['punkt', 'stopwords']
    for package in nltk_packages:
        try:
            nltk.data.find('tokenizers/' + package)
        except LookupError:
            nltk.download(package)

# Call the function to check and download the necessary NLTK packages
check_nltk_packages()

# Function to load and read a text file
def load_and_read_txt(file):
    return file.getvalue().decode('utf-8')

def load_and_read_docx(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
        temp.write(file.getbuffer())
        temp_path = temp.name

    doc = Document(temp_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Delete the temporary file
    os.remove(temp_path)

    return text

# Function to load and read an Excel file
def load_and_read_xlsx(file):
    wb = load_workbook(filename=file)
    ws = wb.active
    data = ws.values
    cols = next(data)
    data_rows = list(data)
    df = pd.DataFrame(data_rows, columns=cols)
    return ' '.join(df.to_string(index=False).split('\n'))

# Function to summarize a report
def summarize_report(report):
    stopWords = set(stopwords.words("english"))
    words = word_tokenize(report)

    freqTable = dict()
    for word in words:
        word = word.lower()
        if word in stopWords:
            continue
        if word in freqTable:
            freqTable[word] += 1
        else:
            freqTable[word] = 1

    sentences = sent_tokenize(report)
    sentenceValue = dict()

    for sentence in sentences:
        for word, freq in freqTable.items():
            if word in sentence.lower():
                if sentence in sentenceValue:
                    sentenceValue[sentence] += freq
                else:
                    sentenceValue[sentence] = freq

    sumValues = 0
    for sentence in sentenceValue:
        sumValues += sentenceValue[sentence]

    # Average value of a sentence from the original report
    average = int(sumValues / len(sentenceValue))

    summary = ''
    for sentence in sentences:
        if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * average)):
            summary += " " + sentence
    return summary

# Function to write a summary to a text file
def write_to_txt(summary):
    with open('summary.txt', 'w') as f:
        f.write(summary)

# Function to write a summary to a Word file
def write_to_docx(summary):
    doc = Document()
    doc.add_paragraph(summary)
    with open('summary.docx', 'wb') as f:
        doc.save(f)

# Function to write a summary to an Excel file
def write_to_xlsx(summary):
    df = pd.DataFrame({'Summary': [summary]})
    df.to_excel('summary.xlsx', index=False)

# Streamlit app
st.title('Doctor\'s Report Summarizer')
st.markdown("This App reads in or analyze an existing patient's report written by the doctor, then write it into "
            "a more easily readable format, without the heavy technical medical or health-related terminologies") 
st.markdown("Please upload a doctor's report in .txt, .docx, or .xlsx format to get started.")

file = st.file_uploader('Upload a doctor\'s report', type=['txt', 'docx', 'xlsx'])

if file is not None:
    if file.type == 'text/plain':
        report = load_and_read_txt(file)
        summary = summarize_report(report)
        print(summary)  # Add this line
        write_to_txt(summary)
    elif file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        report = load_and_read_docx(file)
        summary = summarize_report(report)
        print(summary)  # Add this line
        write_to_docx(summary)
    elif file.type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        report = load_and_read_xlsx(file)
        summary = summarize_report(report)
        print(summary)  # Add this line
        write_to_xlsx(summary)
    else:
        st.write('Unsupported file type')

    st.write(summary)